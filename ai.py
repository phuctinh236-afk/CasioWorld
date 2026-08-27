# ============================================
# File: ai.py
# Code AI VIP - trợ lý lập trình tự học (TF-IDF, không dùng API AI ngoài)
# + Đăng nhập Google, hệ thống VIP (quota ảnh/file, tên 7 màu, avatar, sửa ảnh)
# + Sidebar: nút mở/đóng menu, "Chat mới", danh sách lịch sử chat, bấm vào xem lại
# ============================================

import os
import io
import re
import json
import math
import uuid
import base64
import secrets
import zipfile
import sqlite3
import datetime
import requests
from flask import (
    Flask, request, jsonify, render_template_string,
    session, redirect, url_for
)
from werkzeug.middleware.proxy_fix import ProxyFix
from authlib.integrations.flask_client import OAuth
from PIL import Image, ImageFilter, ImageEnhance
from pypdf import PdfReader
import docx as docx_lib


# ---------- TF-IDF + cosine (pure Python — không cần sklearn / numpy) ----------
_TOKEN_RE = re.compile(
    r"[a-zA-Z0-9àáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđ_]+",
    re.I,
)


def _tokenize(text):
    return _TOKEN_RE.findall((text or "").lower())


class SimpleTfidf:
    """TF-IDF vectorizer thuần Python (thay sklearn.TfidfVectorizer)."""

    def __init__(self):
        self.vocab = {}  # term -> index
        self.idf = []
        self.n_docs = 0

    def fit_transform(self, docs):
        docs_tokens = [_tokenize(d) for d in docs]
        self.n_docs = max(len(docs_tokens), 1)
        df = {}
        for tokens in docs_tokens:
            for t in set(tokens):
                df[t] = df.get(t, 0) + 1
        terms = sorted(df.keys())
        self.vocab = {t: i for i, t in enumerate(terms)}
        self.idf = [0.0] * len(terms)
        for t, i in self.vocab.items():
            # giống sklearn smooth_idf: log((1+N)/(1+df)) + 1
            self.idf[i] = math.log((1 + self.n_docs) / (1 + df[t])) + 1.0
        return [self._tfidf_vec(tokens) for tokens in docs_tokens]

    def transform(self, docs):
        return [self._tfidf_vec(_tokenize(d)) for d in docs]

    def _tfidf_vec(self, tokens):
        if not self.vocab:
            return []
        n = len(self.vocab)
        tf = [0.0] * n
        if not tokens:
            return tf
        for t in tokens:
            i = self.vocab.get(t)
            if i is not None:
                tf[i] += 1.0
        total = float(len(tokens)) or 1.0
        return [(tf[i] / total) * self.idf[i] if tf[i] else 0.0 for i in range(n)]


def _cosine_similarity_row(query_vec, matrix):
    """Trả về list cosine(query, từng vector trong matrix)."""

    def _norm(v):
        return math.sqrt(sum(x * x for x in v)) or 1e-12

    qn = _norm(query_vec)
    scores = []
    for row in matrix:
        dot = sum(a * b for a, b in zip(query_vec, row))
        scores.append(dot / (qn * _norm(row)))
    return scores


# ---------- Cấu hình ----------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
AVATAR_DIR = os.path.join(BASE_DIR, "static", "avatars")
UPLOAD_DIR = os.path.join(BASE_DIR, "static", "uploads")
DB_FILE = os.path.join(BASE_DIR, "app.db")
KB_FILE = os.path.join(BASE_DIR, "knowledge_base.json")

MAX_IMAGES_FREE_PER_DAY = 3
MAX_FILES_FREE_PER_DAY = 3

os.makedirs(AVATAR_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)

app = Flask(__name__, static_folder="static")
app.secret_key = os.environ.get("SECRET_KEY", "doi-key-nay-truoc-khi-deploy-that")
app.config["MAX_CONTENT_LENGTH"] = 25 * 1024 * 1024

# Render (và hầu hết PaaS) chạy Flask sau 1 reverse proxy: nếu không khai báo
# ProxyFix thì url_for(..., _external=True) sẽ trả về "http://" thay vì
# "https://" -> redirect_uri gửi lên Google KHÔNG khớp với URI đã đăng ký trên
# Google Cloud Console -> lỗi "redirect_uri_mismatch". Đây là nguyên nhân phổ
# biến nhất khiến đăng nhập Google báo lỗi trên môi trường deploy thật.
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)

# Cookie session cần "Secure" (chỉ gửi qua HTTPS) + SameSite=Lax để cookie
# vẫn còn sau khi trình duyệt được Google redirect ngược lại (nếu không,
# session bị mất giữa bước /login/google và /auth/google/callback, khiến
# state/nonce không khớp -> lỗi "mismatching_state" hoặc "CSRF Warning").
app.config["SESSION_COOKIE_SECURE"] = os.environ.get("FLASK_ENV") != "development"
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["SESSION_COOKIE_HTTPONLY"] = True

GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET", "")
GOOGLE_LOGIN_CONFIGURED = bool(GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET)

oauth = OAuth(app)
google = oauth.register(
    name="google",
    client_id=GOOGLE_CLIENT_ID or "missing-client-id",
    client_secret=GOOGLE_CLIENT_SECRET or "missing-client-secret",
    server_metadata_url="https://accounts.google.com/.well-known/openid-configuration",
    client_kwargs={"scope": "openid email profile"},
)


# ---------- Database ----------
def get_db():
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_db()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            google_id TEXT UNIQUE NOT NULL,
            email TEXT,
            name TEXT,
            avatar_url TEXT,
            is_vip INTEGER DEFAULT 0,
            vip_expire TEXT,
            image_count INTEGER DEFAULT 0,
            file_count INTEGER DEFAULT 0,
            last_reset TEXT
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS conversations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            owner_key TEXT NOT NULL,
            title TEXT,
            created_at TEXT
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id INTEGER NOT NULL,
            role TEXT NOT NULL,
            content TEXT NOT NULL,
            created_at TEXT
        )
    """)
    conn.commit()
    conn.close()


init_db()


def today_str():
    return datetime.date.today().isoformat()


def now_str():
    return datetime.datetime.now().isoformat()


def get_user_by_id(user_id):
    conn = get_db()
    row = conn.execute("SELECT * FROM users WHERE id=?", (user_id,)).fetchone()
    conn.close()
    return row


def get_or_create_user(google_id, email, name, picture):
    conn = get_db()
    row = conn.execute("SELECT * FROM users WHERE google_id=?", (google_id,)).fetchone()
    if row is None:
        conn.execute(
            "INSERT INTO users (google_id, email, name, avatar_url, last_reset) VALUES (?,?,?,?,?)",
            (google_id, email, name, picture, today_str()),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM users WHERE google_id=?", (google_id,)).fetchone()
    conn.close()
    return row


def reset_quota_if_new_day(user_row):
    if user_row["last_reset"] != today_str():
        conn = get_db()
        conn.execute(
            "UPDATE users SET image_count=0, file_count=0, last_reset=? WHERE id=?",
            (today_str(), user_row["id"]),
        )
        conn.commit()
        conn.close()
        return get_user_by_id(user_row["id"])
    return user_row


def is_vip_active(user_row):
    if not user_row["is_vip"]:
        return False
    if user_row["vip_expire"]:
        try:
            expire = datetime.date.fromisoformat(user_row["vip_expire"])
            return expire >= datetime.date.today()
        except ValueError:
            return True
    return True


def current_user():
    uid = session.get("user_id")
    if not uid:
        return None
    row = get_user_by_id(uid)
    if row is None:
        return None
    return reset_quota_if_new_day(row)


def owner_key():
    """Định danh chủ sở hữu hội thoại: user đăng nhập -> 'u:<id>', khách -> 'g:<uuid ngẫu nhiên>' lưu trong session."""
    user = current_user()
    if user is not None:
        return f"u:{user['id']}"
    if "guest_id" not in session:
        session["guest_id"] = uuid.uuid4().hex
    return f"g:{session['guest_id']}"


# ---------- AI lập trình ----------
class CodeAIVIP:
    def __init__(self):
        self.default_kb = [
            {"question": "bạn tên gì", "answer": "Tôi là Code AI VIP, trợ lý lập trình của riêng bạn."},
            {"question": "bạn làm được gì", "answer": "Tôi giải thích khái niệm lập trình, gỡ lỗi code, và học thêm từ bạn."},
            {"question": "tạm biệt", "answer": "Tạm biệt! Hẹn gặp lại."},

            # --- Cơ bản lập trình ---
            {"question": "biến là gì", "answer": "Biến là nơi lưu trữ dữ liệu, có tên gọi, giá trị có thể thay đổi trong lúc chương trình chạy."},
            {"question": "hằng số là gì", "answer": "Hằng số là giá trị được gán 1 lần và không thay đổi trong suốt chương trình."},
            {"question": "hàm là gì", "answer": "Hàm là khối lệnh được đặt tên, nhận tham số và có thể trả về giá trị, giúp tái sử dụng code."},
            {"question": "tham số và đối số khác nhau thế nào", "answer": "Tham số (parameter) là biến khai báo trong định nghĩa hàm, đối số (argument) là giá trị thực tế truyền vào khi gọi hàm."},
            {"question": "vòng lặp for và while khác nhau thế nào", "answer": "For dùng khi biết trước số lần lặp hoặc duyệt tập hợp; while dùng khi lặp theo điều kiện chưa biết trước số lần."},
            {"question": "câu lệnh if else dùng để làm gì", "answer": "If/else dùng để rẽ nhánh: thực hiện 1 khối lệnh nếu điều kiện đúng, khối khác nếu sai."},
            {"question": "mảng là gì", "answer": "Mảng (array) là tập hợp các phần tử cùng kiểu, lưu liên tiếp, truy cập qua chỉ số (index)."},
            {"question": "recursion là gì", "answer": "Đệ quy (recursion) là kỹ thuật hàm tự gọi lại chính nó để giải bài toán nhỏ hơn, cần có điều kiện dừng (base case)."},
            {"question": "kiểu dữ liệu là gì", "answer": "Kiểu dữ liệu (data type) mô tả loại giá trị 1 biến có thể lưu: số nguyên, số thực, chuỗi, boolean, mảng, object..."},

            # --- Cấu trúc dữ liệu & thuật toán ---
            {"question": "list và tuple khác nhau thế nào", "answer": "List có thể thay đổi (mutable) dùng [], tuple không thể thay đổi (immutable) dùng ()."},
            {"question": "stack là gì", "answer": "Stack là cấu trúc dữ liệu LIFO (vào sau ra trước), thường dùng cho undo, đệ quy, duyệt cây."},
            {"question": "queue là gì", "answer": "Queue là cấu trúc dữ liệu FIFO (vào trước ra trước), dùng cho hàng đợi tác vụ, BFS."},
            {"question": "hash table là gì", "answer": "Hash table (bảng băm) lưu dữ liệu dạng key-value, tra cứu trung bình O(1) nhờ hàm băm."},
            {"question": "linked list là gì", "answer": "Linked list là danh sách liên kết, mỗi phần tử (node) chứa dữ liệu và con trỏ tới phần tử tiếp theo."},
            {"question": "big o là gì", "answer": "Big O mô tả độ phức tạp thời gian/không gian của thuật toán khi dữ liệu đầu vào tăng, ví dụ O(1), O(n), O(log n), O(n^2)."},
            {"question": "sắp xếp nhanh quicksort là gì", "answer": "Quicksort là thuật toán sắp xếp chia để trị, chọn phần tử chốt (pivot) rồi phân hoạch mảng, độ phức tạp trung bình O(n log n)."},
            {"question": "binary search là gì", "answer": "Tìm kiếm nhị phân dùng cho mảng đã sắp xếp, mỗi bước loại bỏ 1 nửa vùng tìm kiếm, độ phức tạp O(log n)."},

            # --- OOP ---
            {"question": "lập trình hướng đối tượng là gì", "answer": "OOP (Object-Oriented Programming) tổ chức code quanh các đối tượng (object) chứa dữ liệu và hành vi, dựa trên 4 tính chất: đóng gói, kế thừa, đa hình, trừu tượng."},
            {"question": "class và object khác nhau thế nào", "answer": "Class là khuôn mẫu định nghĩa thuộc tính/phương thức, object là thực thể cụ thể được tạo ra từ class."},
            {"question": "kế thừa trong oop là gì", "answer": "Kế thừa (inheritance) cho phép 1 class con dùng lại thuộc tính/phương thức của class cha, tránh lặp code."},
            {"question": "đa hình trong oop là gì", "answer": "Đa hình (polymorphism) cho phép các class khác nhau dùng chung 1 interface nhưng có cách xử lý riêng."},
            {"question": "interface là gì", "answer": "Interface định nghĩa tập hợp phương thức mà class phải hiện thực, không chứa logic cụ thể."},

            # --- Web / Backend ---
            {"question": "flask là gì", "answer": "Flask là web framework nhẹ của Python để dựng web server/API nhanh."},
            {"question": "django là gì", "answer": "Django là web framework Python đầy đủ tính năng (ORM, admin, auth có sẵn), phù hợp dự án lớn."},
            {"question": "api là gì", "answer": "API (Application Programming Interface) là giao diện cho phép các chương trình giao tiếp với nhau."},
            {"question": "rest api là gì", "answer": "REST API là kiểu thiết kế API dùng HTTP (GET, POST, PUT, DELETE) thao tác trên tài nguyên (resource)."},
            {"question": "http status code 404 nghĩa là gì", "answer": "404 Not Found nghĩa là server không tìm thấy tài nguyên được yêu cầu."},
            {"question": "http status code 500 nghĩa là gì", "answer": "500 Internal Server Error nghĩa là server gặp lỗi nội bộ khi xử lý yêu cầu."},
            {"question": "cookie và session khác nhau thế nào", "answer": "Cookie lưu ở trình duyệt người dùng; session lưu ở server, thường dùng cookie chứa session ID để định danh."},
            {"question": "cors là gì", "answer": "CORS (Cross-Origin Resource Sharing) là cơ chế trình duyệt kiểm soát việc 1 trang web gọi API từ domain khác."},
            {"question": "middleware là gì", "answer": "Middleware là đoạn code chạy giữa request và response, dùng để xử lý auth, logging, validate..."},

            # --- Frontend ---
            {"question": "javascript là gì", "answer": "JavaScript là ngôn ngữ lập trình chạy chủ yếu trên trình duyệt, dùng để làm web tương tác."},
            {"question": "html là gì", "answer": "HTML là ngôn ngữ đánh dấu dùng để xây dựng cấu trúc nội dung trang web."},
            {"question": "css là gì", "answer": "CSS dùng để định dạng, tạo kiểu (màu sắc, bố cục) cho trang web."},
            {"question": "dom là gì", "answer": "DOM (Document Object Model) là cấu trúc cây đại diện cho nội dung HTML, JavaScript dùng DOM để thay đổi trang web."},
            {"question": "react là gì", "answer": "React là thư viện JavaScript để xây dựng giao diện người dùng theo component, phát triển bởi Meta."},
            {"question": "async await trong javascript là gì", "answer": "Async/await là cú pháp giúp viết code bất đồng bộ (promise) trông giống code đồng bộ, dễ đọc hơn."},

            # --- Database ---
            {"question": "sql là gì", "answer": "SQL (Structured Query Language) là ngôn ngữ truy vấn dùng để thao tác với cơ sở dữ liệu quan hệ."},
            {"question": "primary key là gì", "answer": "Primary key là cột (hoặc tổ hợp cột) định danh duy nhất mỗi dòng trong 1 bảng."},
            {"question": "foreign key là gì", "answer": "Foreign key là cột tham chiếu tới primary key của bảng khác, dùng để liên kết dữ liệu giữa các bảng."},
            {"question": "sql và nosql khác nhau thế nào", "answer": "SQL dùng bảng có cấu trúc cố định (quan hệ); NoSQL linh hoạt hơn, lưu dạng document/key-value/graph, vd MongoDB, Redis."},
            {"question": "index trong database là gì", "answer": "Index giúp tăng tốc độ truy vấn bằng cách tạo cấu trúc tra cứu nhanh trên 1 hoặc nhiều cột."},

            # --- Git / công cụ ---
            {"question": "git là gì", "answer": "Git là hệ thống quản lý phiên bản, theo dõi thay đổi code qua commit, push, pull."},
            {"question": "git commit là gì", "answer": "git commit lưu lại 1 điểm chốt thay đổi trong lịch sử code kèm mô tả (message)."},
            {"question": "git branch là gì", "answer": "git branch tạo 1 nhánh phát triển riêng, không ảnh hưởng nhánh chính (main/master) cho tới khi merge."},
            {"question": "git merge và git rebase khác nhau thế nào", "answer": "Merge giữ nguyên lịch sử 2 nhánh và tạo commit gộp; rebase viết lại lịch sử bằng cách di chuyển commit lên đầu nhánh đích, lịch sử thẳng hơn."},
            {"question": "pull request là gì", "answer": "Pull request (PR) là yêu cầu gộp thay đổi từ 1 nhánh/fork vào nhánh chính, thường kèm review code."},
            {"question": "docker là gì", "answer": "Docker đóng gói ứng dụng cùng môi trường chạy vào container, đảm bảo chạy giống nhau ở mọi nơi."},

            # --- Debug / chất lượng code ---
            {"question": "cách debug code", "answer": "Dùng print/log kiểm tra giá trị biến, dùng breakpoint trong IDE, hoặc đọc kỹ traceback để tìm nguyên nhân lỗi."},
            {"question": "exception là gì", "answer": "Exception là lỗi xảy ra lúc chạy chương trình, có thể bắt bằng try/except (Python) hoặc try/catch (JS/Java)."},
            {"question": "unit test là gì", "answer": "Unit test kiểm thử 1 đơn vị code nhỏ (thường là 1 hàm) độc lập, đảm bảo nó hoạt động đúng như kỳ vọng."},
            {"question": "clean code là gì", "answer": "Clean code là code dễ đọc, dễ hiểu, đặt tên rõ ràng, hàm ngắn gọn, tránh lặp code (DRY)."},
            {"question": "refactor là gì", "answer": "Refactor là việc cải tổ cấu trúc code mà không thay đổi hành vi bên ngoài, giúp code sạch và dễ bảo trì hơn."},

            # --- Bảo mật cơ bản ---
            {"question": "sql injection là gì", "answer": "SQL injection là lỗ hổng khi input người dùng bị chèn thẳng vào câu truy vấn SQL, cách phòng là dùng parameterized query."},
            {"question": "xss là gì", "answer": "XSS (Cross-Site Scripting) là lỗ hổng cho phép chèn script độc hại vào trang web hiển thị cho người dùng khác."},
            {"question": "mã hoá mật khẩu như thế nào cho an toàn", "answer": "Không lưu mật khẩu dạng thô, nên băm (hash) bằng thuật toán như bcrypt kèm salt trước khi lưu vào database."},

            # --- Python cụ thể ---
            {"question": "python là gì", "answer": "Python là ngôn ngữ lập trình bậc cao, cú pháp rõ ràng, dùng nhiều trong web, AI, khoa học dữ liệu."},
            {"question": "list comprehension trong python là gì", "answer": "List comprehension là cú pháp ngắn gọn tạo list mới từ 1 iterable, ví dụ [x*2 for x in range(5)]."},
            {"question": "dict trong python là gì", "answer": "Dict (dictionary) là cấu trúc lưu dữ liệu dạng key-value, truy cập giá trị qua key thay vì chỉ số."},
            {"question": "decorator trong python là gì", "answer": "Decorator là hàm bọc 1 hàm khác để thêm chức năng mà không sửa code gốc, dùng cú pháp @ten_decorator."},
            {"question": "self trong python nghĩa là gì", "answer": "self là tham số đại diện cho chính instance đang gọi phương thức trong 1 class Python."},

            # --- Deploy ---
            {"question": "cách deploy web lên render", "answer": "Đẩy code lên GitHub, tạo Web Service mới trên Render trỏ vào repo, cấu hình build/start command (vd gunicorn app:app), Render tự build và chạy."},
            {"question": "biến môi trường env variable là gì", "answer": "Biến môi trường là giá trị cấu hình (API key, mật khẩu...) lưu ngoài code, đọc lúc chạy để tránh lộ thông tin nhạy cảm trong source."},
        ]
        self.knowledge_base = self._load_kb()
        self._rebuild_vectors()

    def _load_kb(self):
        if os.path.exists(KB_FILE):
            try:
                with open(KB_FILE, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, list) and data:
                        return data
            except Exception:
                pass
        return list(self.default_kb)

    def _save_kb(self):
        try:
            with open(KB_FILE, "w", encoding="utf-8") as f:
                json.dump(self.knowledge_base, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    def _rebuild_vectors(self):
        self.questions = [i["question"] for i in self.knowledge_base]
        self.answers = [i["answer"] for i in self.knowledge_base]
        self.vectorizer = SimpleTfidf()
        if self.questions:
            self.question_vectors = self.vectorizer.fit_transform(self.questions)
        else:
            self.question_vectors = []

    def train(self, pairs):
        for q, a in pairs:
            self.knowledge_base.append({"question": q, "answer": a})
        self._rebuild_vectors()
        self._save_kb()

    def respond(self, text):
        if not text.strip():
            return "Vui lòng nhập câu hỏi."
        if not self.question_vectors:
            return "Tôi chưa hiểu. Dạy tôi bằng: học|câu hỏi|câu trả lời"
        vec = self.vectorizer.transform([text])[0]
        sims = _cosine_similarity_row(vec, self.question_vectors)
        idx = max(range(len(sims)), key=lambda i: sims[i])
        if sims[idx] > 0.2:
            return self.answers[idx]
        return "Tôi chưa hiểu. Dạy tôi bằng: học|câu hỏi|câu trả lời"

    def learn_from_input(self, text):
        parts = text.split("|")
        if len(parts) == 3 and parts[0].strip().lower() == "học":
            q, a = parts[1].strip(), parts[2].strip()
            if q and a:
                self.train([(q, a)])
                return "Đã học xong, cảm ơn bạn!"
        return None


ai_vip = CodeAIVIP()


# ---------- Gemini API (free tier) - tuỳ chọn, tự động dùng nếu có key ----------
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")
GEMINI_URL = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent"

SYSTEM_PROMPT = (
    "Bạn là Code AI VIP, một trợ lý lập trình nói tiếng Việt, trả lời ngắn gọn, "
    "chính xác, có ví dụ code khi cần thiết."
)


def call_gemini(user_text, extra_context=""):
    """Gọi Gemini API (free tier). Trả về None nếu chưa cấu hình key hoặc lỗi,
    để hàm gọi có thể tự rơi về (fallback) bot TF-IDF nội bộ."""
    if not GEMINI_API_KEY:
        return None
    prompt = user_text if not extra_context else f"{extra_context}\n\n---\nCâu hỏi: {user_text}"
    try:
        resp = requests.post(
            GEMINI_URL,
            params={"key": GEMINI_API_KEY},
            json={
                "system_instruction": {"parts": [{"text": SYSTEM_PROMPT}]},
                "contents": [{"role": "user", "parts": [{"text": prompt}]}],
            },
            timeout=25,
        )
        if resp.status_code != 200:
            return None
        data = resp.json()
        return data["candidates"][0]["content"]["parts"][0]["text"]
    except Exception:
        return None


CODE_SYSTEM_PROMPT = (
    "Bạn là Code AI VIP, một lập trình viên AI. Khi được yêu cầu viết code, hãy "
    "trả lời bằng code hoàn chỉnh, chạy được, đặt trong khối markdown dạng "
    "```ten_ngon_ngu\\n...code...\\n```. Nếu code gồm nhiều file, hãy dùng nhiều "
    "khối code, mỗi khối bắt đầu bằng dòng comment ghi tên file, ví dụ "
    "'# file: app.py'. Giải thích ngắn gọn bằng tiếng Việt trước/sau code, "
    "không dài dòng."
)


def call_gemini_code(user_text):
    """Gọi Gemini với system prompt chuyên viết code, trả về None nếu chưa có key."""
    if not GEMINI_API_KEY:
        return None
    try:
        resp = requests.post(
            GEMINI_URL,
            params={"key": GEMINI_API_KEY},
            json={
                "system_instruction": {"parts": [{"text": CODE_SYSTEM_PROMPT}]},
                "contents": [{"role": "user", "parts": [{"text": user_text}]}],
            },
            timeout=45,
        )
        if resp.status_code != 200:
            return None
        data = resp.json()
        return data["candidates"][0]["content"]["parts"][0]["text"]
    except Exception:
        return None


# ---------- Gemini API tạo/sửa ảnh bằng AI (free tier) ----------
GEMINI_IMAGE_MODEL = os.environ.get("GEMINI_IMAGE_MODEL", "gemini-2.5-flash-image")
GEMINI_IMAGE_URL = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_IMAGE_MODEL}:generateContent"


def call_gemini_image(prompt, input_image_bytes=None, input_mime="image/png"):
    """Tạo ảnh mới từ mô tả, hoặc sửa ảnh có sẵn theo mô tả (nếu truyền input_image_bytes).
    Trả về (png_bytes, error)."""
    if not GEMINI_API_KEY:
        return None, "Server chưa cấu hình GEMINI_API_KEY nên chưa dùng được tính năng AI tạo/sửa ảnh."
    parts = []
    if input_image_bytes:
        parts.append({
            "inline_data": {
                "mime_type": input_mime,
                "data": base64.b64encode(input_image_bytes).decode("ascii"),
            }
        })
    parts.append({"text": prompt or "Tạo một bức ảnh đẹp, sáng tạo."})
    try:
        resp = requests.post(
            GEMINI_IMAGE_URL,
            params={"key": GEMINI_API_KEY},
            json={
                "contents": [{"role": "user", "parts": parts}],
                "generationConfig": {"responseModalities": ["IMAGE", "TEXT"]},
            },
            timeout=60,
        )
        if resp.status_code != 200:
            return None, f"Gemini lỗi ({resp.status_code}): {resp.text[:300]}"
        data = resp.json()
        candidates = data.get("candidates") or []
        if not candidates:
            return None, "Gemini không trả về ảnh nào."
        for part in candidates[0].get("content", {}).get("parts", []):
            inline = part.get("inlineData") or part.get("inline_data")
            if inline and inline.get("data"):
                return base64.b64decode(inline["data"]), None
        return None, "Gemini không trả về dữ liệu ảnh (có thể model chưa hỗ trợ ảnh ở khu vực này)."
    except Exception as e:
        return None, f"Lỗi khi gọi Gemini: {e}"


# ---------- Đọc nội dung file (free, không cần AI ngoài) ----------
def extract_text_from_file(file_storage):
    """Trả về (text, error). Hỗ trợ txt/md, csv, pdf, docx."""
    filename = file_storage.filename or ""
    ext = os.path.splitext(filename)[1].lower()
    try:
        if ext in (".txt", ".md", ".csv", ".json", ".py", ".js", ".html", ".css"):
            raw = file_storage.read()
            text = raw.decode("utf-8", errors="ignore")
            return text, None
        elif ext == ".pdf":
            reader = PdfReader(file_storage.stream)
            parts = []
            for page in reader.pages[:30]:  # giới hạn 30 trang để tránh quá tải
                parts.append(page.extract_text() or "")
            return "\n".join(parts), None
        elif ext == ".docx":
            document = docx_lib.Document(file_storage.stream)
            parts = [p.text for p in document.paragraphs]
            return "\n".join(parts), None
        else:
            return None, f"Chưa hỗ trợ đọc file loại {ext or '(không rõ)'}."
    except Exception as e:
        return None, f"Không đọc được file: {e}"


# ---------- Auth ----------
@app.route("/login/google")
def login_google():
    if not GOOGLE_LOGIN_CONFIGURED:
        return (
            "Đăng nhập Google chưa được cấu hình trên server này: thiếu biến môi "
            "trường GOOGLE_CLIENT_ID / GOOGLE_CLIENT_SECRET. Vào Google Cloud "
            "Console tạo OAuth Client ID (loại Web application), thêm redirect "
            "URI là <code>%s</code>, rồi khai báo 2 biến môi trường đó trên server."
            % url_for("auth_google_callback", _external=True),
            500,
        )
    # Sinh nonce riêng cho lượt đăng nhập này và lưu vào session để đối chiếu
    # lại ở bước callback (bắt buộc với luồng OpenID Connect, thiếu nonce là
    # một trong các lý do phổ biến khiến việc lấy userinfo/id_token thất bại).
    nonce = secrets.token_urlsafe(24)
    session["google_oauth_nonce"] = nonce
    redirect_uri = url_for("auth_google_callback", _external=True)
    return google.authorize_redirect(redirect_uri, nonce=nonce)


@app.route("/auth/google/callback")
def auth_google_callback():
    if not GOOGLE_LOGIN_CONFIGURED:
        return redirect(url_for("index"))
    try:
        token = google.authorize_access_token()
    except Exception as e:
        # Nguyên nhân thường gặp: session cookie bị mất giữa 2 bước (xem
        # SESSION_COOKIE_SECURE/SAMESITE ở trên), hoặc redirect_uri đăng ký
        # trên Google Cloud Console không khớp domain thật đang chạy.
        return (
            "Đăng nhập Google thất bại (%s). Kiểm tra: 1) redirect URI trên "
            "Google Cloud Console phải khớp CHÍNH XÁC với "
            "%s (kể cả http/https); "
            "2) site đang chạy dưới HTTPS; 3) cookie của trình duyệt không bị "
            "chặn." % (e, url_for("auth_google_callback", _external=True)),
            400,
        )

    userinfo = token.get("userinfo")
    if not userinfo:
        nonce = session.get("google_oauth_nonce")
        try:
            userinfo = google.parse_id_token(token, nonce=nonce)
        except Exception:
            userinfo = None
    if not userinfo:
        # Phương án dự phòng cuối: gọi thẳng userinfo endpoint bằng access
        # token vừa nhận được, không phụ thuộc việc parse id_token.
        try:
            resp = google.get("https://openidconnect.googleapis.com/v1/userinfo", token=token)
            userinfo = resp.json()
        except Exception:
            userinfo = None

    session.pop("google_oauth_nonce", None)

    if not userinfo or "sub" not in userinfo:
        return "Không lấy được thông tin tài khoản Google, vui lòng thử đăng nhập lại.", 400

    google_id = userinfo["sub"]
    email = userinfo.get("email", "")
    name = userinfo.get("name", email or "User")
    picture = userinfo.get("picture", "")
    user = get_or_create_user(google_id, email, name, picture)
    session["user_id"] = user["id"]
    return redirect(url_for("index"))


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("index"))


def user_public_dict(user_row):
    if user_row is None:
        return None
    return {
        "name": user_row["name"],
        "email": user_row["email"],
        "avatar_url": user_row["avatar_url"] or "",
        "is_vip": bool(is_vip_active(user_row)),
        "image_count": user_row["image_count"],
        "file_count": user_row["file_count"],
        "image_limit": MAX_IMAGES_FREE_PER_DAY,
        "file_limit": MAX_FILES_FREE_PER_DAY,
    }


@app.route("/me")
def me():
    return jsonify(user_public_dict(current_user()))


@app.route("/premium-info")
def premium_info():
    user = current_user()
    return jsonify({
        "is_vip": bool(user and is_vip_active(user)),
        "benefits": [
            "Gửi ảnh không giới hạn (tài khoản thường: 3 ảnh/ngày)",
            "Gửi file không giới hạn (tài khoản thường: 3 file/ngày)",
            "Tên hiển thị 7 màu + icon VIP kế bên tên",
            "Đổi avatar tuỳ chọn",
            "Sửa ảnh bằng bộ lọc (grayscale, làm mờ, làm nét, xoay, chỉnh sáng)",
            "Tạo ảnh AI & sửa ảnh AI theo mô tả — miễn phí cho mọi tài khoản đã đăng nhập",
            "Đọc file, viết code, nén code thành .zip — miễn phí cho mọi tài khoản đã đăng nhập",
        ]
    })


# ---------- Hội thoại (sidebar / lịch sử chat) ----------
def make_title_from_message(text):
    text = text.strip().replace("\n", " ")
    if len(text) > 40:
        text = text[:40].rstrip() + "..."
    return text or "Cuộc trò chuyện mới"


@app.route("/conversations")
def list_conversations():
    ok = owner_key()
    conn = get_db()
    rows = conn.execute(
        "SELECT id, title, created_at FROM conversations WHERE owner_key=? ORDER BY id DESC LIMIT 100",
        (ok,),
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/conversations/new", methods=["POST"])
def new_conversation():
    ok = owner_key()
    conn = get_db()
    cur = conn.execute(
        "INSERT INTO conversations (owner_key, title, created_at) VALUES (?,?,?)",
        (ok, "Cuộc trò chuyện mới", now_str()),
    )
    conn.commit()
    conv_id = cur.lastrowid
    conn.close()
    return jsonify({"id": conv_id, "title": "Cuộc trò chuyện mới"})


@app.route("/conversations/<int:conv_id>/messages")
def get_messages(conv_id):
    ok = owner_key()
    conn = get_db()
    conv = conn.execute(
        "SELECT * FROM conversations WHERE id=? AND owner_key=?", (conv_id, ok)
    ).fetchone()
    if conv is None:
        conn.close()
        return jsonify({"error": "Không tìm thấy cuộc trò chuyện."}), 404
    rows = conn.execute(
        "SELECT role, content FROM messages WHERE conversation_id=? ORDER BY id ASC",
        (conv_id,),
    ).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])


@app.route("/conversations/<int:conv_id>", methods=["DELETE"])
def delete_conversation(conv_id):
    ok = owner_key()
    conn = get_db()
    conv = conn.execute(
        "SELECT * FROM conversations WHERE id=? AND owner_key=?", (conv_id, ok)
    ).fetchone()
    if conv is None:
        conn.close()
        return jsonify({"error": "Không tìm thấy cuộc trò chuyện."}), 404
    conn.execute("DELETE FROM messages WHERE conversation_id=?", (conv_id,))
    conn.execute("DELETE FROM conversations WHERE id=?", (conv_id,))
    conn.commit()
    conn.close()
    return jsonify({"ok": True})


# ---------- Chat (giờ có lưu lịch sử) ----------
@app.route("/")
def index():
    return render_template_string(HTML_TEMPLATE)


@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json(silent=True) or {}
    msg = data.get("message", "")
    conv_id = data.get("conversation_id")
    ok = owner_key()

    conn = get_db()
    if conv_id:
        conv = conn.execute(
            "SELECT * FROM conversations WHERE id=? AND owner_key=?", (conv_id, ok)
        ).fetchone()
        if conv is None:
            conn.close()
            return jsonify({"error": "Cuộc trò chuyện không hợp lệ."}), 404
    else:
        cur = conn.execute(
            "INSERT INTO conversations (owner_key, title, created_at) VALUES (?,?,?)",
            (ok, make_title_from_message(msg), now_str()),
        )
        conn.commit()
        conv_id = cur.lastrowid

    CODE_KEYWORDS = ("viết code", "viết hàm", "viết chương trình", "code cho",
                     "sửa code", "sửa lỗi code", "debug", "viết class", "viết api")
    looks_like_code_request = any(k in msg.lower() for k in CODE_KEYWORDS)

    learned = ai_vip.learn_from_input(msg)
    if learned:
        reply = learned
    else:
        # ưu tiên Gemini (free tier) nếu đã cấu hình GEMINI_API_KEY; dùng
        # system prompt chuyên viết code khi câu hỏi rõ ràng là yêu cầu code,
        # để trả lời kèm khối ```code``` sạch (frontend dùng khối này để
        # bật nút "Tải .zip").
        if looks_like_code_request:
            reply = call_gemini_code(msg)
        else:
            reply = call_gemini(msg)
        if reply is None:
            reply = ai_vip.respond(msg)  # fallback bot TF-IDF nội bộ khi chưa có API key

    conn.execute(
        "INSERT INTO messages (conversation_id, role, content, created_at) VALUES (?,?,?,?)",
        (conv_id, "user", msg, now_str()),
    )
    conn.execute(
        "INSERT INTO messages (conversation_id, role, content, created_at) VALUES (?,?,?,?)",
        (conv_id, "ai", reply, now_str()),
    )
    conn.commit()
    conn.close()

    return jsonify({"reply": reply, "conversation_id": conv_id})


# ---------- Upload ảnh / file (có quota) ----------
def check_and_consume_quota(user, kind):
    if user is None:
        return False, "Bạn cần đăng nhập bằng Google để gửi ảnh/file."
    if is_vip_active(user):
        return True, None
    field = "image_count" if kind == "image" else "file_count"
    limit = MAX_IMAGES_FREE_PER_DAY if kind == "image" else MAX_FILES_FREE_PER_DAY
    if user[field] >= limit:
        loai = "ảnh" if kind == "image" else "file"
        return False, (f"Bạn đã gửi {limit} {loai} hôm nay (giới hạn tài khoản thường). "
                        f"Nâng cấp VIP để gửi không giới hạn, hoặc quay lại vào ngày mai.")
    conn = get_db()
    conn.execute(f"UPDATE users SET {field} = {field} + 1 WHERE id=?", (user["id"],))
    conn.commit()
    conn.close()
    return True, None


@app.route("/upload/image", methods=["POST"])
def upload_image():
    user = current_user()
    ok, msg = check_and_consume_quota(user, "image")
    if not ok:
        return jsonify({"error": msg}), 403
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "Thiếu file ảnh."}), 400
    filename = f"{user['id']}_{int(datetime.datetime.now().timestamp())}_{file.filename}"
    path = os.path.join(UPLOAD_DIR, filename)
    file.save(path)
    return jsonify({"ok": True, "url": f"/static/uploads/{filename}"})


@app.route("/upload/file", methods=["POST"])
def upload_file():
    user = current_user()
    ok, msg = check_and_consume_quota(user, "file")
    if not ok:
        return jsonify({"error": msg}), 403
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "Thiếu file."}), 400
    filename = f"{user['id']}_{int(datetime.datetime.now().timestamp())}_{file.filename}"
    path = os.path.join(UPLOAD_DIR, filename)
    file.save(path)
    return jsonify({"ok": True, "url": f"/static/uploads/{filename}"})


# ---------- Sửa ảnh (Pillow, MIỄN PHÍ - chỉ cần đăng nhập) ----------
@app.route("/edit-image", methods=["POST"])
def edit_image():
    user = current_user()
    if user is None:
        return jsonify({"error": "Bạn cần đăng nhập Google để dùng tính năng này."}), 403
    file = request.files.get("file")
    action = request.form.get("action", "grayscale")
    if not file:
        return jsonify({"error": "Thiếu file ảnh."}), 400
    img = Image.open(file.stream).convert("RGB")
    if action == "grayscale":
        img = img.convert("L")
    elif action == "blur":
        img = img.filter(ImageFilter.GaussianBlur(4))
    elif action == "sharpen":
        img = img.filter(ImageFilter.SHARPEN)
    elif action == "brighten":
        img = ImageEnhance.Brightness(img).enhance(1.4)
    elif action == "darken":
        img = ImageEnhance.Brightness(img).enhance(0.7)
    elif action == "rotate":
        img = img.rotate(90, expand=True)
    else:
        return jsonify({"error": "Hành động chỉnh sửa không hợp lệ."}), 400
    filename = f"edit_{user['id']}_{int(datetime.datetime.now().timestamp())}.png"
    path = os.path.join(UPLOAD_DIR, filename)
    img.save(path)
    return jsonify({"ok": True, "url": f"/static/uploads/{filename}"})


# ---------- Đọc file (free, không cần AI ngoài) ----------
@app.route("/read-file", methods=["POST"])
def read_file():
    user = current_user()
    if user is None:
        return jsonify({"error": "Bạn cần đăng nhập Google để dùng tính năng này."}), 403
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "Thiếu file."}), 400

    text, err = extract_text_from_file(file)
    if err:
        return jsonify({"error": err}), 400

    preview = text[:6000]  # giới hạn độ dài để không quá tải khung chat / Gemini

    ask_ai = request.form.get("ask_ai", "true") == "true"
    summary = None
    if ask_ai:
        summary = call_gemini(
            "Tóm tắt ngắn gọn nội dung file dưới đây bằng tiếng Việt, "
            "nêu các ý chính.",
            extra_context=preview,
        )
    return jsonify({
        "ok": True,
        "content_preview": preview,
        "truncated": len(text) > 6000,
        "summary": summary,  # None nếu chưa cấu hình Gemini hoặc lỗi
    })


# ---------- Sửa file (free, không cần AI ngoài) ----------
@app.route("/edit-file", methods=["POST"])
def edit_file_route():
    user = current_user()
    if user is None:
        return jsonify({"error": "Bạn cần đăng nhập Google để dùng tính năng này."}), 403
    file = request.files.get("file")
    find = request.form.get("find", "")
    replace = request.form.get("replace", "")
    if not file:
        return jsonify({"error": "Thiếu file."}), 400
    if not find:
        return jsonify({"error": "Cần nhập nội dung cần tìm (find)."}), 400

    filename = file.filename or ""
    ext = os.path.splitext(filename)[1].lower()

    try:
        if ext in (".txt", ".md", ".csv", ".json", ".py", ".js", ".html", ".css"):
            raw = file.read().decode("utf-8", errors="ignore")
            edited = raw.replace(find, replace)
            out_name = f"edit_{user['id']}_{int(datetime.datetime.now().timestamp())}{ext}"
            path = os.path.join(UPLOAD_DIR, out_name)
            with open(path, "w", encoding="utf-8") as f:
                f.write(edited)
            return jsonify({"ok": True, "url": f"/static/uploads/{out_name}"})

        elif ext == ".docx":
            document = docx_lib.Document(file.stream)
            for p in document.paragraphs:
                if find in p.text:
                    for run in p.runs:
                        if find in run.text:
                            run.text = run.text.replace(find, replace)
            out_name = f"edit_{user['id']}_{int(datetime.datetime.now().timestamp())}.docx"
            path = os.path.join(UPLOAD_DIR, out_name)
            document.save(path)
            return jsonify({"ok": True, "url": f"/static/uploads/{out_name}"})

        else:
            return jsonify({"error": f"Chưa hỗ trợ sửa file loại {ext or '(không rõ)'} (chỉ hỗ trợ txt/md/csv/json/docx...)."}), 400
    except Exception as e:
        return jsonify({"error": f"Lỗi khi sửa file: {e}"}), 500


# ---------- Tạo ảnh bằng AI (FREE cho mọi tài khoản đã đăng nhập) ----------
@app.route("/generate/image", methods=["POST"])
def generate_image():
    user = current_user()
    # Không còn giới hạn riêng cho VIP nữa: mọi tài khoản Google đã đăng nhập
    # đều dùng được tính năng tạo ảnh AI. Vẫn tính vào quota ảnh/ngày chung để
    # tránh 1 tài khoản gọi API key miễn phí quá nhiều gây khoá key.
    ok, msg = check_and_consume_quota(user, "image")
    if not ok:
        return jsonify({"error": msg}), 403
    data = request.get_json(silent=True) or {}
    prompt = (data.get("prompt") or "").strip()
    if not prompt:
        return jsonify({"error": "Thiếu mô tả ảnh muốn tạo."}), 400
    img_bytes, err = call_gemini_image(prompt)
    if err:
        return jsonify({"error": err}), 502
    filename = f"gen_{user['id']}_{int(datetime.datetime.now().timestamp())}.png"
    path = os.path.join(UPLOAD_DIR, filename)
    with open(path, "wb") as f:
        f.write(img_bytes)
    return jsonify({"ok": True, "url": f"/static/uploads/{filename}"})


# ---------- Sửa ảnh bằng AI theo mô tả (FREE, khác với /edit-image dùng bộ lọc Pillow) ----------
@app.route("/edit-image-ai", methods=["POST"])
def edit_image_ai():
    user = current_user()
    if user is None:
        return jsonify({"error": "Bạn cần đăng nhập Google để dùng tính năng này."}), 403
    ok, msg = check_and_consume_quota(user, "image")
    if not ok:
        return jsonify({"error": msg}), 403
    file = request.files.get("file")
    prompt = (request.form.get("prompt") or "").strip()
    if not file:
        return jsonify({"error": "Thiếu file ảnh."}), 400
    if not prompt:
        return jsonify({"error": "Thiếu mô tả cách muốn sửa ảnh."}), 400
    img_bytes = file.read()
    out_bytes, err = call_gemini_image(prompt, input_image_bytes=img_bytes,
                                        input_mime=file.mimetype or "image/png")
    if err:
        return jsonify({"error": err}), 502
    filename = f"aiedit_{user['id']}_{int(datetime.datetime.now().timestamp())}.png"
    path = os.path.join(UPLOAD_DIR, filename)
    with open(path, "wb") as f:
        f.write(out_bytes)
    return jsonify({"ok": True, "url": f"/static/uploads/{filename}"})


# ---------- Nén code AI viết ra thành file .zip (FREE) ----------
@app.route("/export-zip", methods=["POST"])
def export_zip():
    user = current_user()
    if user is None:
        return jsonify({"error": "Bạn cần đăng nhập Google để tải file .zip."}), 403
    data = request.get_json(silent=True) or {}
    files = data.get("files") or []
    if not isinstance(files, list) or not files:
        return jsonify({"error": "Không có file code nào để nén."}), 400

    zip_name = f"code_{user['id']}_{int(datetime.datetime.now().timestamp())}.zip"
    zip_path = os.path.join(UPLOAD_DIR, zip_name)
    used_names = set()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, item in enumerate(files):
            name = (item.get("filename") or f"file_{i+1}.txt").strip().lstrip("/\\")
            name = name.replace("..", "_") or f"file_{i+1}.txt"
            base_name = name
            n = 1
            while name in used_names:
                stem, dot, ext = base_name.rpartition(".")
                name = f"{stem or base_name}_{n}{(dot + ext) if dot else ''}"
                n += 1
            used_names.add(name)
            content = item.get("content") or ""
            zf.writestr(name, content)

    return jsonify({"ok": True, "url": f"/static/uploads/{zip_name}"})


# ---------- Video / audio: chưa có dịch vụ tạo sinh miễn phí để kết nối ----------
@app.route("/generate/video", methods=["POST"])
def generate_video():
    return jsonify({"error": "Tính năng tạo video bằng AI chưa được kết nối dịch vụ tạo sinh."}), 501


@app.route("/generate/audio", methods=["POST"])
def generate_audio():
    return jsonify({"error": "Tính năng tạo audio bằng AI chưa được kết nối dịch vụ tạo sinh."}), 501


# ---------- Avatar tuỳ chỉnh (chỉ VIP) ----------
@app.route("/profile/avatar", methods=["POST"])
def change_avatar():
    user = current_user()
    if user is None or not is_vip_active(user):
        return jsonify({"error": "Chỉ VIP mới được đổi avatar tuỳ chọn."}), 403
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "Thiếu file ảnh."}), 400
    ext = os.path.splitext(file.filename)[1] or ".png"
    filename = f"avatar_{user['id']}{ext}"
    path = os.path.join(AVATAR_DIR, filename)
    file.save(path)
    url = f"/static/avatars/{filename}"
    conn = get_db()
    conn.execute("UPDATE users SET avatar_url=? WHERE id=?", (url, user["id"]))
    conn.commit()
    conn.close()
    return jsonify({"ok": True, "url": url})


# ---------- Kích hoạt VIP thủ công (chưa có cổng thanh toán thật) ----------
def set_vip(user_id, days=30):
    expire = (datetime.date.today() + datetime.timedelta(days=days)).isoformat()
    conn = get_db()
    conn.execute("UPDATE users SET is_vip=1, vip_expire=? WHERE id=?", (expire, user_id))
    conn.commit()
    conn.close()


@app.route("/admin/activate-vip", methods=["POST"])
def admin_activate_vip():
    data = request.get_json(silent=True) or {}
    if data.get("admin_key") != os.environ.get("ADMIN_KEY", "change-me"):
        return jsonify({"error": "Không có quyền."}), 403
    set_vip(data.get("user_id"), int(data.get("days", 30)))
    return jsonify({"ok": True})


# ---------- Giao diện ----------
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="vi">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Code AI VIP</title>
<style>
  :root {
    --bg:#1a1a19; --panel:#232322; --sidebar:#181817; --bubble-user:#3a3a38; --bubble-ai:#2b2b29;
    --accent:#d97757; --text:#ece9e6; --muted:#9a968f; --border:#333;
  }
  * { box-sizing:border-box; }
  html, body { height:100%; }
  body { margin:0; font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;
    background:var(--bg); color:var(--text);
    height:100vh; height:100dvh; /* 100dvh né thanh địa chỉ/tab của trình duyệt di động */
    display:flex; overflow:hidden; }

  /* ---- Sidebar ---- */
  #sidebar {
    width:270px; background:var(--sidebar); border-right:1px solid var(--border);
    display:flex; flex-direction:column; flex-shrink:0;
    position:fixed; top:0; bottom:0; left:0; z-index:50;
    transform:translateX(-100%); transition:transform 0.2s ease;
  }
  #sidebar.open { transform:translateX(0); }
  #overlay {
    position:fixed; inset:0; background:rgba(0,0,0,0.5); z-index:40;
    display:none;
  }
  #overlay.show { display:block; }
  .sidebar-top { padding:16px; display:flex; align-items:center; justify-content:space-between; }
  .sidebar-top .brand-row { display:flex; align-items:center; gap:8px; }
  .sidebar-top .brand { font-weight:700; font-size:16px; }
  .logo-img { width:26px; height:26px; border-radius:8px; object-fit:cover; flex-shrink:0; }
  .new-chat-btn {
    margin:0 16px 8px; padding:10px 12px; background:var(--panel); border:1px solid var(--border);
    border-radius:10px; color:var(--text); display:flex; align-items:center; gap:8px;
    cursor:pointer; font-size:14px;
  }
  .nav-item {
    margin:0 16px; padding:9px 12px; border-radius:10px; color:var(--text); display:flex;
    align-items:center; gap:8px; cursor:pointer; font-size:13.5px;
  }
  .nav-item:hover { background:var(--panel); }
  .nav-divider { border-top:1px solid var(--border); margin:10px 0; }
  .sidebar-section-title { padding:6px 16px; font-size:11px; color:var(--muted); text-transform:uppercase; letter-spacing:0.05em; }
  #conv-list { flex:1; overflow-y:auto; padding:0 8px; min-height:40px; }
  .conv-item {
    padding:10px 10px; border-radius:8px; font-size:13.5px; color:var(--text);
    cursor:pointer; white-space:nowrap; overflow:hidden; text-overflow:ellipsis;
    display:flex; align-items:center; justify-content:space-between; gap:6px;
  }
  .conv-item:hover { background:var(--panel); }
  .conv-item.active { background:var(--bubble-user); }
  .conv-item .del-btn { opacity:0; background:none; border:none; color:var(--muted); cursor:pointer; font-size:13px; }
  .conv-item:hover .del-btn { opacity:1; }
  .conv-empty { padding:10px; font-size:12.5px; color:var(--muted); }

  .sidebar-bottom {
    border-top:1px solid var(--border); padding:12px 16px;
    padding-bottom:calc(12px + env(safe-area-inset-bottom, 0px));
    display:flex; align-items:center; gap:10px;
  }
  .sidebar-bottom .avatar-img, .sidebar-bottom .avatar-fallback {
    width:34px; height:34px; border-radius:50%; object-fit:cover; background:var(--accent);
    flex-shrink:0; display:flex; align-items:center; justify-content:center; color:#1a1a19; font-weight:700; font-size:13px;
  }
  .sidebar-bottom .user-info { flex:1; min-width:0; display:flex; flex-direction:column; }
  .sidebar-bottom .user-info .line1 { display:flex; align-items:center; gap:4px; overflow:hidden; }
  .sidebar-bottom .user-info .email { font-size:11px; color:var(--muted); overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }
  .sidebar-bottom .logout-btn { background:none; border:none; color:var(--muted); cursor:pointer; font-size:13px; padding:4px; }
  .login-btn { background:var(--accent); color:#1a1a19; border:none; padding:9px 12px; border-radius:8px;
    font-size:13px; font-weight:600; cursor:pointer; text-decoration:none; width:100%; text-align:center; display:block; }

  /* ---- Main ---- */
  #main { flex:1; display:flex; flex-direction:column; margin-left:0; height:100%; min-width:0; }
  header { padding:12px 16px; border-bottom:1px solid var(--border); display:flex; align-items:center;
    gap:10px; flex-shrink:0; }
  .brand-row { display:flex; align-items:center; gap:10px; }
  .hamburger { background:none; border:none; color:var(--text); font-size:20px; cursor:pointer; padding:4px 6px; }
  header h1 { font-size:15px; margin:0; font-weight:600; }
  header span.sub { font-size:11.5px; color:var(--muted); margin-left:4px; display:none; }
  @media (min-width:520px) { header span.sub { display:inline; } }

  .rainbow-name { font-weight:700; font-size:13px;
    background:linear-gradient(90deg,#ff4d4d,#ffa64d,#ffe14d,#4dff88,#4dd2ff,#4d79ff,#c14dff);
    background-size:300% 100%; -webkit-background-clip:text; background-clip:text; color:transparent;
    animation:rainbow-move 4s linear infinite; white-space:nowrap; overflow:hidden; text-overflow:ellipsis; }
  @keyframes rainbow-move { 0%{background-position:0% 50%} 100%{background-position:300% 50%} }
  .vip-badge-icon { width:15px; height:15px; flex-shrink:0; }
  .plain-name { font-size:13px; color:var(--text); white-space:nowrap; overflow:hidden; text-overflow:ellipsis; }

  /* ---- Chat ---- */
  #chat-box { flex:1; overflow-y:auto; padding:20px; max-width:760px; width:100%; margin:0 auto; min-height:0; }
  .row { display:flex; margin-bottom:16px; }
  .row.user { justify-content:flex-end; }
  .row.ai { justify-content:flex-start; }
  .bubble { max-width:78%; padding:12px 16px; border-radius:14px; line-height:1.5; font-size:14.5px; white-space:pre-wrap; }
  .row.user .bubble { background:var(--bubble-user); border-bottom-right-radius:4px; }
  .row.ai .bubble { background:var(--bubble-ai); border:1px solid var(--border); border-bottom-left-radius:4px; }
  .avatar { width:26px; height:26px; border-radius:6px; background:var(--accent); color:#1a1a19;
    display:flex; align-items:center; justify-content:center; font-size:12px; font-weight:bold; margin-right:10px; flex-shrink:0; overflow:hidden; }
  .avatar img { width:100%; height:100%; object-fit:cover; }
  .thinking-bubble { color:var(--muted); font-style:italic; display:flex; align-items:center; gap:8px; }
  .think-dot { width:7px; height:7px; border-radius:50%; background:var(--accent); flex-shrink:0; animation:think-pulse 1s ease-in-out infinite; }
  @keyframes think-pulse { 0%,100%{opacity:0.3;transform:scale(0.8);} 50%{opacity:1;transform:scale(1.2);} }

  #input-area {
    border-top:1px solid var(--border); padding:12px 16px;
    padding-bottom:calc(18px + env(safe-area-inset-bottom, 14px)); /* né thanh tab/gesture bar */
    flex-shrink:0;
  }
  #input-inner { max-width:760px; margin:0 auto; display:flex; align-items:flex-end; gap:8px;
    background:var(--panel); border:1px solid var(--border); border-radius:16px; padding:10px 12px; }
  #user-input { flex:1; resize:none; border:none; outline:none; background:transparent; color:var(--text);
    font-size:14.5px; max-height:140px; font-family:inherit; }
  .icon-btn { background:transparent; border:none; color:var(--muted); cursor:pointer; font-size:16px; }
  button#send-btn { background:var(--accent); border:none; color:#1a1a19; font-weight:600; padding:8px 16px;
    border-radius:10px; cursor:pointer; font-size:13.5px; }
  .hint { text-align:center; color:var(--muted); font-size:11px; margin-top:8px; max-width:760px; margin-left:auto; margin-right:auto; }

  /* ---- Modal (Dự án / VIP Premium) ---- */
  #modal-overlay { position:fixed; inset:0; background:rgba(0,0,0,0.6); z-index:100; display:none;
    align-items:center; justify-content:center; padding:20px; }
  #modal-overlay.show { display:flex; }
  .modal-box { background:var(--panel); border:1px solid var(--border); border-radius:16px; padding:20px;
    max-width:400px; width:100%; max-height:80vh; overflow-y:auto; }
  .modal-box h2 { margin:0 0 12px; font-size:17px; display:flex; align-items:center; gap:8px; }
  .modal-box img.modal-logo { width:22px; height:22px; }
  .modal-box ul { margin:0; padding-left:18px; font-size:13.5px; line-height:1.7; color:var(--text); }
  .modal-box p { font-size:13px; color:var(--muted); line-height:1.6; }
  .modal-close { margin-top:16px; background:var(--accent); color:#1a1a19; border:none; padding:9px 14px;
    border-radius:8px; font-weight:600; font-size:13px; cursor:pointer; width:100%; }
</style>
</head>
<body>

<div id="overlay" onclick="closeSidebar()"></div>

<div id="sidebar">
  <div class="sidebar-top">
    <div class="brand-row">
      <img class="logo-img" src="/static/assets/aiviplogo.png" onerror="this.style.display='none'">
      <div class="brand">Code AI VIP</div>
    </div>
    <button class="hamburger" onclick="closeSidebar()">✕</button>
  </div>
  <div class="new-chat-btn" onclick="startNewChat()">＋ Chat mới</div>
  <div class="nav-item" onclick="openModal('projects')">📁 Dự án</div>
  <div class="nav-item" onclick="openModal('premium')">
    <img class="vip-badge-icon" src="/static/assets/aiviplogo.png" onerror="this.style.display='none'"> AI VIP Premium
  </div>
  <div class="nav-divider"></div>
  <div class="sidebar-section-title">Gần đây</div>
  <div id="conv-list"></div>
  <div class="sidebar-bottom" id="sidebar-user-box"></div>
</div>

<div id="modal-overlay" onclick="if(event.target===this) closeModal()">
  <div class="modal-box" id="modal-box"></div>
</div>

<div id="main">
  <header>
    <button class="hamburger" onclick="openSidebar()">☰</button>
    <img class="logo-img" src="/static/assets/aiviplogo.png" onerror="this.style.display='none'">
    <h1>Code AI VIP <span class="sub">trợ lý lập trình tự học</span></h1>
  </header>

  <div id="chat-box"></div>

  <div id="input-area">
    <div id="input-inner">
      <button class="icon-btn" title="Gửi/đọc ảnh & file" onclick="document.getElementById('file-input').click()">📎</button>
      <input type="file" id="file-input" style="display:none" onchange="uploadPickedFile()">
      <button class="icon-btn" title="Sửa ảnh/file" onclick="document.getElementById('edit-file-input').click()">🛠</button>
      <input type="file" id="edit-file-input" style="display:none" onchange="editFilePicked()">
      <button class="icon-btn" title="Tạo ảnh bằng AI" onclick="openGenerateImageModal()">🎨</button>
      <textarea id="user-input" rows="1" placeholder="Nhắn cho Code AI VIP..."></textarea>
      <button id="send-btn" onclick="sendMessage()">Gửi</button>
    </div>
    <div class="hint" id="quota-hint">Đang tải thông tin tài khoản...</div>
  </div>
</div>

<script>
const chatBox = document.getElementById('chat-box');
const input = document.getElementById('user-input');
const sendBtn = document.getElementById('send-btn');
const sidebarUserBox = document.getElementById('sidebar-user-box');
const quotaHint = document.getElementById('quota-hint');
const sidebar = document.getElementById('sidebar');
const overlay = document.getElementById('overlay');
const convList = document.getElementById('conv-list');
const modalOverlay = document.getElementById('modal-overlay');
const modalBox = document.getElementById('modal-box');

let ME = null;
let CURRENT_CONV_ID = null;

function openSidebar() { sidebar.classList.add('open'); overlay.classList.add('show'); loadConversations(); }
function closeSidebar() { sidebar.classList.remove('open'); overlay.classList.remove('show'); }

function closeModal() { modalOverlay.classList.remove('show'); modalBox.innerHTML = ''; }

async function openModal(kind) {
  if (kind === 'projects') {
    modalBox.innerHTML = `
      <h2>📁 Dự án</h2>
      <p>Tính năng nhóm các cuộc trò chuyện theo dự án đang được phát triển, sẽ có sớm. Hiện tại bạn có thể dùng "Chat mới" và xem lại lịch sử ở mục Gần đây.</p>
      <button class="modal-close" onclick="closeModal()">Đóng</button>
    `;
  } else if (kind === 'premium') {
    const res = await fetch('/premium-info');
    const data = await res.json();
    const items = data.benefits.map(b => `<li>${b}</li>`).join('');
    const status = data.is_vip
      ? '<p style="color:#f5c542;font-weight:600;">Tài khoản của bạn đang là VIP ✨</p>'
      : '<p>Bạn đang dùng tài khoản thường. Liên hệ quản trị viên để nâng cấp VIP.</p>';
    modalBox.innerHTML = `
      <h2><img class="modal-logo" src="/static/assets/aiviplogo.png" onerror="this.style.display='none'"> AI VIP Premium</h2>
      ${status}
      <ul>${items}</ul>
      <button class="modal-close" onclick="closeModal()">Đóng</button>
    `;
  }
  modalOverlay.classList.add('show');
}

const CODE_BLOCK_RE = /```([a-zA-Z0-9_+-]*)\\n([\\s\\S]*?)```/g;

function renderAiContent(bubble, text) {
  let lastIndex = 0, match, hasCode = false;
  const codeFiles = [];
  CODE_BLOCK_RE.lastIndex = 0;
  while ((match = CODE_BLOCK_RE.exec(text)) !== null) {
    hasCode = true;
    const before = text.slice(lastIndex, match.index).trim();
    if (before) {
      const p = document.createElement('div');
      p.textContent = before; p.style.marginBottom = '8px';
      bubble.appendChild(p);
    }
    const lang = (match[1] || 'txt').toLowerCase();
    let code = match[2];
    let filename = null;
    const fileLineMatch = code.match(/^\\s*(?:#|\\/\\/)\\s*file:\\s*(\\S+)\\s*\\n/i);
    if (fileLineMatch) {
      filename = fileLineMatch[1];
      code = code.slice(fileLineMatch[0].length);
    } else {
      const extMap = {python:'py', py:'py', javascript:'js', js:'js', html:'html', css:'css', json:'json', bash:'sh', sh:'sh', java:'java', c:'c', cpp:'cpp'};
      filename = `code_${codeFiles.length + 1}.${extMap[lang] || 'txt'}`;
    }
    codeFiles.push({ filename, content: code });
    const pre = document.createElement('pre');
    pre.style.cssText = 'background:#1a1a19;border:1px solid var(--border);border-radius:8px;padding:10px;overflow-x:auto;font-size:12.5px;margin:8px 0;';
    const codeEl = document.createElement('code');
    codeEl.textContent = code;
    pre.appendChild(codeEl);
    bubble.appendChild(pre);
    lastIndex = CODE_BLOCK_RE.lastIndex;
  }
  const rest = text.slice(lastIndex);
  if (!hasCode) {
    const p = document.createElement('div');
    p.textContent = text;
    bubble.appendChild(p);
  } else if (rest.trim()) {
    const p = document.createElement('div');
    p.textContent = rest.trim();
    bubble.appendChild(p);
  }
  if (hasCode) {
    const zipBtn = document.createElement('button');
    zipBtn.textContent = '⬇️ Tải code (.zip)';
    zipBtn.className = 'modal-close';
    zipBtn.style.cssText = 'width:auto;margin-top:8px;padding:7px 12px;font-size:12.5px;';
    zipBtn.onclick = () => downloadCodeZip(codeFiles);
    bubble.appendChild(zipBtn);
  }
}

async function downloadCodeZip(files) {
  try {
    const res = await fetch('/export-zip', {
      method: 'POST', headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ files })
    });
    const data = await res.json();
    if (data.error) { addBubble('⚠️ ' + data.error, 'ai'); return; }
    window.open(data.url, '_blank');
  } catch (e) {
    addBubble('⚠️ Không nén được file .zip.', 'ai');
  }
}

function addBubble(text, who) {
  const row = document.createElement('div');
  row.className = 'row ' + who;
  if (who === 'ai') {
    const avatar = document.createElement('div');
    avatar.className = 'avatar'; avatar.textContent = 'C';
    row.appendChild(avatar);
  }
  const bubble = document.createElement('div');
  bubble.className = 'bubble';
  if (who === 'ai') {
    renderAiContent(bubble, text);
  } else {
    bubble.textContent = text;
  }
  row.appendChild(bubble);
  chatBox.appendChild(row);
  chatBox.scrollTop = chatBox.scrollHeight;
}

function addImageBubble(url, who) {
  const row = document.createElement('div');
  row.className = 'row ' + who;
  if (who === 'ai') {
    const avatar = document.createElement('div');
    avatar.className = 'avatar'; avatar.textContent = 'C';
    row.appendChild(avatar);
  }
  const bubble = document.createElement('div');
  bubble.className = 'bubble';
  const img = document.createElement('img');
  img.src = url;
  img.style.cssText = 'max-width:100%;border-radius:10px;display:block;';
  bubble.appendChild(img);
  row.appendChild(bubble);
  chatBox.appendChild(row);
  chatBox.scrollTop = chatBox.scrollHeight;
}

function openGenerateImageModal() {
  modalBox.innerHTML = `
    <h2>🎨 Tạo ảnh bằng AI</h2>
    <p>Mô tả bức ảnh bạn muốn tạo (miễn phí, cần đăng nhập Google).</p>
    <textarea id="gen-image-prompt" rows="3" style="width:100%;margin-bottom:12px;padding:9px;border-radius:8px;border:1px solid #3a3a38;background:#1a1a19;color:#ece9e6;" placeholder="VD: một chú mèo phi hành gia trong vũ trụ, phong cách tranh sơn dầu"></textarea>
    <button class="modal-close" onclick="doGenerateImage()">Tạo ảnh</button>
  `;
  modalOverlay.classList.add('show');
}

async function doGenerateImage() {
  const promptEl = document.getElementById('gen-image-prompt');
  const prompt = promptEl ? promptEl.value.trim() : '';
  if (!prompt) return;
  closeModal();
  addBubble(prompt, 'user');
  const thinkInterval = startThinkingSteps();
  try {
    const res = await fetch('/generate/image', {
      method: 'POST', headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ prompt })
    });
    const data = await res.json();
    stopThinkingSteps(thinkInterval);
    if (data.error) addBubble('⚠️ ' + data.error, 'ai');
    else { addImageBubble(data.url, 'ai'); await loadMe(); }
  } catch (e) {
    stopThinkingSteps(thinkInterval);
    addBubble('Lỗi kết nối tới server.', 'ai');
  }
}

function clearChat() {
  chatBox.innerHTML = '';
  addBubble('Chào bạn! Hỏi mình về lập trình, hoặc dạy mình bằng: học|câu hỏi|câu trả lời', 'ai');
}

async function loadMe() {
  try {
    const res = await fetch('/me');
    ME = await res.json();
  } catch (e) {
    ME = null;
  }
  renderUserBox();
}

function renderUserBox() {
  if (!ME || !ME.name) {
    sidebarUserBox.innerHTML = '<a class="login-btn" href="/login/google">Đăng nhập Google</a>';
    quotaHint.textContent = 'Đăng nhập ở menu ☰ để gửi ảnh/file cho AI.';
    return;
  }
  const nameClass = ME.is_vip ? 'rainbow-name' : 'plain-name';
  const vipIcon = ME.is_vip ? `<img class="vip-badge-icon" src="/static/assets/aiviplogo.png" onerror="this.style.display='none'">` : '';
  const avatarHtml = ME.avatar_url
    ? `<img class="avatar-img" src="${ME.avatar_url}">`
    : `<div class="avatar-fallback">${ME.name.charAt(0).toUpperCase()}</div>`;
  sidebarUserBox.innerHTML = `
    ${avatarHtml}
    <div class="user-info">
      <div class="line1"><span class="${nameClass}">${ME.name}</span>${vipIcon}</div>
      <span class="email">${ME.email || ''}</span>
    </div>
    <a href="/logout" class="logout-btn" title="Đăng xuất">⏻</a>
  `;
  quotaHint.textContent = ME.is_vip
    ? 'Tài khoản VIP: gửi ảnh/file không giới hạn.'
    : `Ảnh: ${ME.image_count}/${ME.image_limit} hôm nay · File: ${ME.file_count}/${ME.file_limit} hôm nay`;
}

async function loadConversations() {
  try {
    const res = await fetch('/conversations');
    const list = await res.json();
    convList.innerHTML = '';
    if (list.length === 0) {
      convList.innerHTML = '<div class="conv-empty">Chưa có cuộc trò chuyện nào.</div>';
      return;
    }
    list.forEach(c => {
      const item = document.createElement('div');
      item.className = 'conv-item' + (c.id === CURRENT_CONV_ID ? ' active' : '');
      item.innerHTML = `<span style="overflow:hidden;text-overflow:ellipsis;">${c.title}</span>
                         <button class="del-btn" onclick="event.stopPropagation();deleteConv(${c.id})">🗑</button>`;
      item.onclick = () => openConversation(c.id);
      convList.appendChild(item);
    });
  } catch (e) {
    convList.innerHTML = '<div class="conv-empty">Không tải được lịch sử chat.</div>';
  }
}

async function openConversation(id) {
  CURRENT_CONV_ID = id;
  const res = await fetch(`/conversations/${id}/messages`);
  const msgs = await res.json();
  chatBox.innerHTML = '';
  msgs.forEach(m => addBubble(m.content, m.role === 'user' ? 'user' : 'ai'));
  closeSidebar();
}

async function deleteConv(id) {
  await fetch(`/conversations/${id}`, { method: 'DELETE' });
  if (id === CURRENT_CONV_ID) { CURRENT_CONV_ID = null; clearChat(); }
  loadConversations();
}

function startNewChat() {
  CURRENT_CONV_ID = null;
  clearChat();
  closeSidebar();
}

const THINK_STEPS = [
  'Đang đọc câu hỏi của bạn...',
  'Đang tra cứu kiến thức đã học...',
  'Đang phân tích yêu cầu lập trình...',
  'Đang soạn câu trả lời...'
];

function addThinkingBubble() {
  const row = document.createElement('div');
  row.className = 'row ai';
  row.id = 'thinking-row';
  const avatar = document.createElement('div');
  avatar.className = 'avatar'; avatar.textContent = 'C';
  row.appendChild(avatar);
  const bubble = document.createElement('div');
  bubble.className = 'bubble thinking-bubble';
  bubble.innerHTML = `<span class="think-dot"></span><span id="think-current">${THINK_STEPS[0]}</span>`;
  row.appendChild(bubble);
  chatBox.appendChild(row);
  chatBox.scrollTop = chatBox.scrollHeight;
}

function startThinkingSteps() {
  addThinkingBubble();
  let i = 0;
  return setInterval(() => {
    i = (i + 1) % THINK_STEPS.length;
    const el = document.getElementById('think-current');
    if (el) el.textContent = THINK_STEPS[i];
  }, 850);
}

function stopThinkingSteps(intervalId) {
  clearInterval(intervalId);
  const row = document.getElementById('thinking-row');
  if (row) row.remove();
}

async function sendMessage() {
  const message = input.value.trim();
  if (!message) return;
  addBubble(message, 'user');
  input.value = ''; input.style.height = 'auto';
  sendBtn.disabled = true;
  const thinkInterval = startThinkingSteps();
  try {
    const response = await fetch('/chat', {
      method: 'POST', headers: {'Content-Type':'application/json'},
      body: JSON.stringify({ message, conversation_id: CURRENT_CONV_ID })
    });
    const data = await response.json();
    stopThinkingSteps(thinkInterval);
    addBubble(data.reply, 'ai');
    if (data.conversation_id) CURRENT_CONV_ID = data.conversation_id;
  } catch (e) {
    stopThinkingSteps(thinkInterval);
    addBubble('Lỗi kết nối tới server.', 'ai');
  } finally {
    sendBtn.disabled = false;
  }
}

async function uploadPickedFile() {
  const fileInput = document.getElementById('file-input');
  const file = fileInput.files[0];
  if (!file) return;
  addBubble('Đã gửi: ' + file.name, 'user');

  const isImage = file.type.startsWith('image/');
  const ext = '.' + (file.name.split('.').pop() || '').toLowerCase();
  const docExts = ['.txt', '.md', '.csv', '.json', '.pdf', '.docx', '.py', '.js', '.html', '.css'];

  const form = new FormData();
  form.append('file', file);

  if (isImage) {
    const res = await fetch('/upload/image', { method: 'POST', body: form });
    const data = await res.json();
    if (data.error) addBubble('⚠️ ' + data.error, 'ai');
    else { addBubble('Đã lưu ảnh.', 'ai'); await loadMe(); }
  } else if (docExts.includes(ext)) {
    const res = await fetch('/read-file', { method: 'POST', body: form });
    const data = await res.json();
    if (data.error) {
      addBubble('⚠️ ' + data.error, 'ai');
    } else {
      let msg = data.summary ? data.summary : (`Nội dung file:
` + data.content_preview);
      if (data.truncated) msg += `

(đã cắt bớt vì file khá dài)`;
      addBubble(msg, 'ai');
    }
  } else {
    const res = await fetch('/upload/file', { method: 'POST', body: form });
    const data = await res.json();
    if (data.error) addBubble('⚠️ ' + data.error, 'ai');
    else { addBubble('Đã lưu file.', 'ai'); await loadMe(); }
  }
  fileInput.value = '';
}

let EDIT_FILE = null;

function editFilePicked() {
  const fileInput = document.getElementById('edit-file-input');
  const file = fileInput.files[0];
  if (!file) return;
  EDIT_FILE = file;
  const isImage = file.type.startsWith('image/');

  if (isImage) {
    const actions = [
      ['grayscale', 'Trắng đen'], ['blur', 'Làm mờ'], ['sharpen', 'Làm nét'],
      ['brighten', 'Tăng sáng'], ['darken', 'Giảm sáng'], ['rotate', 'Xoay 90°']
    ];
    modalBox.innerHTML = `
      <h2>🛠 Sửa ảnh: ${file.name}</h2>
      <p style="margin-top:0;">Bộ lọc nhanh:</p>
      <div style="display:flex;flex-wrap:wrap;gap:8px;margin:12px 0;">
        ${actions.map(([a, label]) => `<button class="modal-close" style="width:auto;flex:1 1 45%;" onclick="doEditImage('${a}')">${label}</button>`).join('')}
      </div>
      <div class="nav-divider"></div>
      <p>Hoặc sửa bằng AI theo mô tả (miễn phí):</p>
      <input id="ai-edit-prompt" placeholder="VD: đổi nền thành bầu trời hoàng hôn" style="width:100%;margin-bottom:10px;padding:9px;border-radius:8px;border:1px solid #3a3a38;background:#1a1a19;color:#ece9e6;">
      <button class="modal-close" onclick="doEditImageAI()">✨ Sửa bằng AI</button>
      <button class="modal-close" style="background:#3a3a38;color:#ece9e6;margin-top:8px;" onclick="closeModal()">Huỷ</button>
    `;
  } else {
    modalBox.innerHTML = `
      <h2>🛠 Sửa file: ${file.name}</h2>
      <p>Tìm & thay thế nội dung (hỗ trợ txt/md/csv/json/docx...):</p>
      <input id="edit-find" placeholder="Tìm..." style="width:100%;margin-bottom:8px;padding:9px;border-radius:8px;border:1px solid #3a3a38;background:#1a1a19;color:#ece9e6;">
      <input id="edit-replace" placeholder="Thay bằng..." style="width:100%;margin-bottom:14px;padding:9px;border-radius:8px;border:1px solid #3a3a38;background:#1a1a19;color:#ece9e6;">
      <button class="modal-close" onclick="doEditFile()">Sửa file</button>
    `;
  }
  modalOverlay.classList.add('show');
  fileInput.value = '';
}

async function doEditImage(action) {
  const form = new FormData();
  form.append('file', EDIT_FILE);
  form.append('action', action);
  const res = await fetch('/edit-image', { method: 'POST', body: form });
  const data = await res.json();
  closeModal();
  if (data.error) addBubble('⚠️ ' + data.error, 'ai');
  else addBubble('Đã sửa ảnh xong: ' + location.origin + data.url, 'ai');
}

async function doEditImageAI() {
  const promptEl = document.getElementById('ai-edit-prompt');
  const prompt = promptEl ? promptEl.value.trim() : '';
  if (!prompt) return;
  const form = new FormData();
  form.append('file', EDIT_FILE);
  form.append('prompt', prompt);
  closeModal();
  addBubble('✨ Sửa ảnh: ' + prompt, 'user');
  const thinkInterval = startThinkingSteps();
  try {
    const res = await fetch('/edit-image-ai', { method: 'POST', body: form });
    const data = await res.json();
    stopThinkingSteps(thinkInterval);
    if (data.error) addBubble('⚠️ ' + data.error, 'ai');
    else { addImageBubble(data.url, 'ai'); await loadMe(); }
  } catch (e) {
    stopThinkingSteps(thinkInterval);
    addBubble('Lỗi kết nối tới server.', 'ai');
  }
}

async function doEditFile() {
  const find = document.getElementById('edit-find').value;
  const replace = document.getElementById('edit-replace').value;
  const form = new FormData();
  form.append('file', EDIT_FILE);
  form.append('find', find);
  form.append('replace', replace);
  const res = await fetch('/edit-file', { method: 'POST', body: form });
  const data = await res.json();
  closeModal();
  if (data.error) addBubble('⚠️ ' + data.error, 'ai');
  else addBubble('Đã sửa file xong: ' + location.origin + data.url, 'ai');
}

input.addEventListener('keydown', function(e) {
  if (e.key === 'Enter' && !e.shiftKey) { e.preventDefault(); sendMessage(); }
});
input.addEventListener('input', function() {
  input.style.height = 'auto';
  input.style.height = Math.min(input.scrollHeight, 140) + 'px';
});

clearChat();
loadMe();
loadConversations();
</script>
</body>
</html>
"""

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)

# ============================================
# HƯỚNG DẪN TRIỂN KHAI
# 1. requirements.txt cần đủ (đã bỏ scikit-learn / numpy / playwright):
#       flask
#       pillow
#       authlib
#       requests
#       gunicorn
#       pypdf
#       python-docx
#       werkzeug
#       cloudscraper
#
# 2. Biến môi trường trên Render:
#       GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET (từ Google Cloud Console,
#         redirect URI: https://<domain>/auth/google/callback)
#       SECRET_KEY (chuỗi ngẫu nhiên dài)
#       ADMIN_KEY (mật khẩu để tự kích hoạt VIP qua /admin/activate-vip)
#
# 3. Lịch sử chat được lưu theo "chủ sở hữu":
#    - Đã đăng nhập Google -> gắn với tài khoản, xem lại được ở bất kỳ máy nào đã đăng nhập.
#    - Chưa đăng nhập (khách) -> gắn với 1 session cookie tạm, mất khi xoá cookie/trình duyệt khác.
#
# 4. Trên Render free tier, ổ đĩa (app.db, ảnh upload) có thể bị xoá khi
#    deploy lại. Muốn giữ lịch sử chat vĩnh viễn, cần gắn Render Disk
#    hoặc chuyển database sang PostgreSQL.
#
# 5. TF-IDF chatbot dùng pure Python (SimpleTfidf) — chạy được trên Termux 2GB RAM.
# ============================================
